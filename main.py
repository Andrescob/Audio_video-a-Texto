import os
from tkinter import Tk, filedialog
from moviepy.editor import AudioFileClip
import speech_recognition as sr
from docx import Document

def seleccionar_archivo():
    # Abrir el explorador de archivos para seleccionar el archivo
    Tk().withdraw()  # Oculta la ventana principal de Tkinter
    archivo = filedialog.askopenfilename(title="Seleccionar archivo de audio o video", filetypes=(("Archivos de Video", "*.mp4;*.mov;*.avi"), ("Archivos de Audio", "*.mp3;*.wav")))
    return archivo

def convertir_audio_a_texto(archivo_audio):
    # Usar SpeechRecognition para convertir el audio a texto
    recognizer = sr.Recognizer()
    with sr.AudioFile(archivo_audio) as source:
        audio = recognizer.record(source)
    try:
        texto = recognizer.recognize_google(audio, language="es-CO")  # Reconocimiento en español (Colombia)
        return texto
    except Exception as e:
        print(f"Error al convertir el audio a texto: {e}")
        return None

def generar_documento(texto):
    # Crear un archivo .doc con el texto convertido
    doc = Document()
    doc.add_heading("Transcripción de Audio/Video", 0)
    doc.add_paragraph(texto)
    doc.save(os.path.expanduser("~/Downloads/Transcripcion.docx"))
    print("Archivo generado en Descargas: Transcripcion.docx")

def convertir_video_a_audio(archivo_video):
    # Extraer audio del video usando moviepy
    clip = AudioFileClip(archivo_video)
    audio_path = "temp_audio.wav"
    clip.write_audiofile(audio_path)
    return audio_path

def main():
    archivo = seleccionar_archivo()
    
    if archivo.lower().endswith(('.mp3', '.wav')):  # Si es un archivo de audio
        archivo_audio = archivo
    elif archivo.lower().endswith(('.mp4', '.mov', '.avi')):  # Si es un archivo de video
        archivo_audio = convertir_video_a_audio(archivo)
    else:
        print("Tipo de archivo no soportado")
        return

    print("Convirtiendo archivo...")
    texto = convertir_audio_a_texto(archivo_audio)
    if texto:
        generar_documento(texto)
    else:
        print("No se pudo convertir el audio a texto")

if __name__ == "__main__":
    main()

